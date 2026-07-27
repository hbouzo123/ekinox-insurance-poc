import os
import json
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# Color palette definition (Ekinox theme)
COLOR_PRIMARY_HEX = "86318E"      # Ekinox Purple
COLOR_SECONDARY_HEX = "0D9488"    # Teal
COLOR_DARK_HEX = "1F2937"         # Dark Charcoal / Off-black
COLOR_LIGHT_HEX = "F9FAFB"        # Very Light Gray
COLOR_MUTED_HEX = "6B7280"        # Medium Gray
COLOR_BORDER_HEX = "E5E7EB"       # Light Gray border

COLOR_PRIMARY = RGBColor(134, 49, 142)
COLOR_SECONDARY = RGBColor(13, 148, 136)
COLOR_DARK = RGBColor(31, 41, 55)
COLOR_MUTED = RGBColor(107, 114, 128)

def set_cell_background(cell, color_hex):
    """Set the background color (shading) of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set the internal margins (padding) of a table cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex, sz="36"):
    """Set a thick left border on a cell (for callout boxes)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    # Left border: thick
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz) # 36 = 4.5pt width
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    # Others: nil
    for border_name in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{border_name}')
        node.set(qn('w:val'), 'nil')
        tcBorders.append(node)
        
    tcPr.append(tcBorders)

def set_table_borders(table, color_hex):
    """Apply clean, horizontal-only borders to a table (no vertical borders)."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Top and bottom of the table
    for border_name in ['top', 'bottom']:
        node = OxmlElement(f'w:{border_name}')
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '10') # ~1.2pt
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color_hex)
        tblBorders.append(node)
        
    # Inside horizontal borders
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '5') # ~0.6pt
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), color_hex)
    tblBorders.append(insideH)
    
    # Disable vertical borders
    for border_name in ['left', 'right', 'insideV']:
        node = OxmlElement(f'w:{border_name}')
        node.set(qn('w:val'), 'nil')
        tblBorders.append(node)
        
    tblPr.append(tblBorders)

def add_heading_1(doc, text):
    """Add a stylized Heading 1."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    return p

def add_heading_2(doc, text):
    """Add a styled Heading 2."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECONDARY
    return p

def add_heading_3(doc, text):
    """Add a styled Heading 3."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK
    return p

def add_paragraph(doc, text, bold_prefix="", indent=0):
    """Add a paragraph with optional bold prefix and indent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
        
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Calibri'
        r_prefix.font.bold = True
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = COLOR_DARK
        
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_bullet_point(doc, text, bold_prefix=""):
    """Add a styled bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Calibri'
        r_prefix.font.bold = True
        r_prefix.font.size = Pt(11)
        r_prefix.font.color.rgb = COLOR_DARK
        
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_DARK
    return p

def add_callout(doc, text, title=""):
    """Create a beautifully shaded callout box with a thick left border."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set full width (6.5 inches)
    table.columns[0].width = Inches(6.5)
    cell = table.rows[0].cells[0]
    cell.width = Inches(6.5)
    
    # Styling
    set_cell_background(cell, COLOR_LIGHT_HEX)
    set_cell_left_border(cell, COLOR_SECONDARY_HEX, sz="36") # Teal thick border
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        r_title = p.add_run(title + "\n")
        r_title.font.name = 'Arial'
        r_title.font.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = COLOR_SECONDARY
        
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_DARK
    
    # Spacer paragraph after table
    doc.add_paragraph().paragraph_format.space_before = Pt(8)

def add_styled_table(doc, headers, data, column_widths=None):
    """Add a beautifully styled table with headers and alternating rows."""
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
        set_cell_background(hdr_cells[idx], COLOR_PRIMARY_HEX)
        set_cell_margins(hdr_cells[idx], top=140, bottom=140, left=150, right=150)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Set data rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        # Alternate background colors
        bg_color = COLOR_LIGHT_HEX if r_idx % 2 == 1 else "FFFFFF"
        
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(0)
            
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_DARK
                
            # Formatting alignments
            if "€" in str(val) or "%" in str(val) or str(val).replace(" ", "").isdigit():
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
    # Apply custom column widths if specified
    if column_widths:
        for idx, width in enumerate(column_widths):
            for row in table.rows:
                row.cells[idx].width = Inches(width)
                
    set_table_borders(table, COLOR_BORDER_HEX)
    
    # Spacer paragraph after table
    doc.add_paragraph().paragraph_format.space_before = Pt(8)

def create_cover_page(doc, title, subtitle, date_str, confidentiality, client_name):
    """Generate a highly polished corporate cover page."""
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.different_first_page_header_footer = True
    
    # Add large top space
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(140)
    
    # Stylized Ekinox Header Block
    p_logo = doc.add_paragraph()
    run_logo = p_logo.add_run("ekinox")
    run_logo.font.name = 'Arial'
    run_logo.font.size = Pt(32)
    run_logo.font.bold = True
    run_logo.font.color.rgb = COLOR_PRIMARY
    
    # Accent line
    p_line = doc.add_paragraph()
    run_line = p_line.add_run("____________________________________________________")
    run_line.font.bold = True
    run_line.font.color.rgb = COLOR_SECONDARY
    
    # Title
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(30)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run(title)
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(100)
    run_sub = p_sub.add_run(subtitle)
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = COLOR_MUTED
    
    # Metadata table on cover page
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    
    metadata = [
        ("Préparé pour :", client_name),
        ("Date de livraison :", date_str),
        ("Statut du document :", confidentiality),
        ("Auteur :", "Ekinox Consulting Team")
    ]
    
    for idx, (label, val) in enumerate(metadata):
        cell_lbl = table.rows[idx].cells[0]
        cell_lbl.text = label
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.runs[0]
        run_lbl.font.bold = True
        run_lbl.font.color.rgb = COLOR_MUTED
        run_lbl.font.size = Pt(10)
        
        cell_val = table.rows[idx].cells[1]
        cell_val.text = val
        p_val = cell_val.paragraphs[0]
        run_val = p_val.runs[0]
        run_val.font.color.rgb = COLOR_DARK
        run_val.font.size = Pt(10)
        if label == "Statut du document :":
            run_val.font.bold = True
            run_val.font.color.rgb = COLOR_PRIMARY
            
    # Set borders and margins to nil for cover page metadata table
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'bottom', 'left', 'right']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'nil')
                tcBorders.append(node)
            tcPr.append(tcBorders)
            set_cell_margins(cell, top=40, bottom=40, left=0, right=0)
            
    # Page Break after Cover Page
    doc.add_page_break()

def setup_headers_footers(doc, doc_title):
    """Configure running headers and footers for the rest of the document."""
    section = doc.sections[0]
    
    # Header
    header = section.header
    p_hdr = header.paragraphs[0]
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr.paragraph_format.space_after = Pt(0)
    run_hdr = p_hdr.add_run(f"Ekinox | {doc_title} — Document Confidentiel")
    run_hdr.font.size = Pt(8.5)
    run_hdr.font.italic = True
    run_hdr.font.color.rgb = COLOR_MUTED
    
    # Footer
    footer = section.footer
    p_ftr = footer.paragraphs[0]
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ftr.paragraph_format.space_after = Pt(0)
    
    run_ftr = p_ftr.add_run("Confidentialité SanlamAllianz - Ekinox  |  Page ")
    run_ftr.font.size = Pt(8.5)
    run_ftr.font.color.rgb = COLOR_MUTED
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    p_ftr._p.append(fldSimple)
    
    run_ftr_mid = p_ftr.add_run(" sur ")
    run_ftr_mid.font.size = Pt(8.5)
    run_ftr_mid.font.color.rgb = COLOR_MUTED
    
    fldSimple2 = OxmlElement('w:fldSimple')
    fldSimple2.set(qn('w:instr'), 'NUMPAGES')
    p_ftr._p.append(fldSimple2)

def build_sales_docx(data):
    """Build the Digital Sales Accelerator word document from JSON data."""
    print("Generating DOCX for Ekinox Digital Sales Accelerator...")
    doc = docx.Document()
    
    # Cover Page
    create_cover_page(
        doc=doc,
        title=data["title"],
        subtitle="Proposition de Transformation & Business Case MVP (8 Capacités)",
        date_str="Juillet 2026",
        confidentiality="STRICTEMENT CONFIDENTIEL — PROPOSITION PARTENAIRE",
        client_name="SanlamAllianz Group"
    )
    
    # Configure Headers & Footers (will appear page 2+)
    setup_headers_footers(doc, "Digital Sales Accelerator")
    
    # 1. Introduction et Positionnement
    add_heading_1(doc, "1. Positionnement Produit : Digital Sales Platform")
    add_paragraph(doc, "SanlamAllianz a besoin d'une nouvelle capacité commerciale omnicanale pour accélérer l'acquisition, la qualification et la conversion de ses prospects. Plutôt que de développer un simple 'chatbot' ou projet IA déconnecté des enjeux d'affaires, Ekinox propose le déploiement d'un accélérateur commercial complet : Ekinox Digital Sales Accelerator.")
    add_paragraph(doc, "Ce produit résout de manière native la dispersion des leads, le taux d'abandon élevé des formulaires en ligne et la perte de productivité des agents commerciaux. Il s'intègre harmonieusement dans l'écosystème d'assurance existant sans dépendances IT lourdes.")
    
    # 2. Product Blueprint
    add_heading_1(doc, "2. Vision Produit & Principes de Conception")
    add_paragraph(doc, data["product_blueprint"]["vision"])
    add_heading_2(doc, "Principes de Conception du MVP")
    for principle in data["product_blueprint"]["principles"]:
        add_paragraph(doc, principle["description"], bold_prefix=f"• {principle['name']} : ")
        
    # 3. Architecture Technique et Intégration Socle Commun
    add_heading_1(doc, "3. Architecture Technique & Socle Commun")
    add_paragraph(doc, data["technical_architecture"]["overview"])
    add_paragraph(doc, data["technical_architecture"]["integration"], bold_prefix="Intégration SI & CRM : ")
    
    add_heading_2(doc, "Composants Technologiques Clés")
    # Parse core components into bullets
    components = data["technical_architecture"]["core_components"].split(". ")
    for comp in components:
        cleaned_comp = comp.strip()
        if not cleaned_comp:
            continue
        # Clean number prefix if present (e.g. "1. " or just "1 ")
        if cleaned_comp[0].isdigit():
            dot_idx = cleaned_comp.find('.')
            if dot_idx != -1 and dot_idx < 3:
                cleaned_comp = cleaned_comp[dot_idx+1:].strip()
            else:
                cleaned_comp = cleaned_comp[1:].strip()
        add_bullet_point(doc, cleaned_comp)
            
    doc.add_page_break()
    
    # 4. Les 8 Capacités Métier de l'Accélérateur V1
    add_heading_1(doc, "4. Cadrage des 8 Capacités du Digital Sales Accelerator")
    add_paragraph(doc, "Chaque capacité de l'accélérateur est présentée avec un double alignement : d'une part la proposition de valeur pour le Client (les gestionnaires de SanlamAllianz et les prospects), et d'autre part les spécifications de construction MVP pour l'Équipe Interne en charge de la solution.")
    
    for idx, cap in enumerate(data["capacities"]):
        # Clean capacity name (remove number if duplicate)
        name = cap["name"]
        if name.startswith(f"{idx+1}."):
            name = name[len(str(idx+1))+2:].strip()
            
        add_heading_2(doc, f"Capacité {idx+1} : {name}")
        
        # Client block
        add_heading_3(doc, "Section Client — Vision Métier & Cas d'Usage")
        add_paragraph(doc, cap["client_section"]["vision"], bold_prefix="Vision & Objectif : ")
        add_paragraph(doc, cap["client_section"]["value_added"], bold_prefix="Valeur Ajoutée Assureur : ")
        add_paragraph(doc, cap["client_section"]["use_case"], bold_prefix="Cas d'Usage Prospect : ")
        add_paragraph(doc, cap["client_section"]["user_impact"], bold_prefix="Impact Utilisateur : ")
        
        # Internal block
        add_heading_3(doc, "Section Équipe Interne — Guide d'Implémentation MVP")
        
        internal_text = (
            f"Données nécessaires pour le MVP :\n{cap['internal_team_section']['data_needed']}\n\n"
            f"Stack technique MVP recommandée :\n{cap['internal_team_section']['tech_stack']}\n\n"
            f"APIs & Services à intégrer/développer :\n{cap['internal_team_section']['apis']}\n\n"
            f"Étapes clés de développement :\n{cap['internal_team_section']['key_steps']}"
        )
        add_callout(doc, internal_text, title=f"Fiche Technique Interne - MVP {name}")
        
    doc.add_page_break()
    
    # 5. Roadmap Produit
    add_heading_1(doc, "5. Trajectoire d'Évolution Produit (Roadmap)")
    add_paragraph(doc, "L'approche d'Ekinox garantit un retour sur investissement rapide avec un MVP ciblé en V1, ouvrant la voie à une trajectoire de maturité produit modulaire et scalable.")
    
    roadmap_headers = ["Phase", "Objectifs & Périmètre Fonctionnel"]
    roadmap_data = []
    # Sort roadmap keys if possible
    rm = data["roadmap"]
    phases_order = ["v1", "v2", "v3", "v4", "v5_v7"]
    for phase_key in phases_order:
        if phase_key in rm:
            phase_label = phase_key.upper().replace("_", "-")
            roadmap_data.append([phase_label, rm[phase_key]])
            
    add_styled_table(doc, roadmap_headers, roadmap_data, column_widths=[1.5, 5.0])
    
    # 6. Business Case & ROI
    add_heading_1(doc, "6. Business Case & Modèle de ROI")
    add_paragraph(doc, "L'évaluation économique du Digital Sales Accelerator repose sur des gains directs d'efficacité commerciale et une optimisation substantielle des budgets marketing.")
    
    # Build structured key value tables for assumptions & impacts
    bc = data["business_case"]
    
    add_heading_2(doc, "Hypothèses de Cadrage et Impacts Attendus")
    add_paragraph(doc, bc["assumptions"], bold_prefix="Hypothèses initiales de volume : ")
    add_paragraph(doc, bc["impact"], bold_prefix="Impacts sur la conversion & CAC : ")
    add_paragraph(doc, bc["roi_calculation"], bold_prefix="Calcul du ROI Financier : ")
    
    add_heading_2(doc, "Seuil de Rentabilité (Point Mort)")
    add_paragraph(doc, bc["break_even"])
    
    # Add a visual representation table for the break-even to wow the client
    be_headers = ["Indicateur", "Valeur Actuelle", "Cible / MVP", "Impact Financier Estimé"]
    be_data = [
        ["Taux de Conversion Lead", "4.0%", "5.5% (+1.5 points)", "+1 500 Ventes / an"],
        ["Coût d'Acquisition Client (CAC)", "50 €", "35 € (-30%)", "+1.5 M€ d'économies marketing"],
        ["Revenus Additionnels Générés", "—", "—", "600 000 € / an"],
        ["Valeur Annuelle Créée", "—", "—", "2 100 000 € / an"]
    ]
    add_styled_table(doc, be_headers, be_data, column_widths=[2.2, 1.3, 1.5, 1.5])
    
    # 7. Conditions de Succès et Gouvernance
    add_heading_1(doc, "7. Modèle de Livraison & Gouvernance")
    add_paragraph(doc, "Le projet est structuré sur un cycle agile de 10 à 12 semaines pour le déploiement du MVP (V1).")
    
    add_heading_2(doc, "Gouvernance Projet")
    add_bullet_point(doc, "Comité de pilotage bimensuel avec le sponsor métier de SanlamAllianz.")
    add_bullet_point(doc, "Ateliers techniques hebdomadaires avec les équipes CRM, Data et Sécurité.")
    add_bullet_point(doc, "Méthodologie Scrum avec des démonstrations fonctionnelles toutes les 2 semaines.")
    
    add_heading_2(doc, "Conditions de Succès Métier")
    add_bullet_point(doc, "Disponibilité rapide d'un accès à l'API WhatsApp Business Sandbox.")
    add_bullet_point(doc, "Accès aux documentations de garanties et exclusions en PDF pour alimenter le Knowledge Hub.")
    add_bullet_point(doc, "Implication active des experts métiers de l'assurance lors de la phase de cadrage du Product Advisor.")
    
    file_path = "c:\\Projects\\Ekinox IA\\Ekinox_Digital_Sales_Accelerator_Business_Case.docx"
    doc.save(file_path)
    print(f"Digital Sales Accelerator DOCX saved to {file_path}")

def build_fraud_docx(data):
    """Build the Fraud Intelligence Accelerator word document from JSON data."""
    print("Generating DOCX for Ekinox Fraud Intelligence Accelerator...")
    doc = docx.Document()
    
    # Cover Page
    create_cover_page(
        doc=doc,
        title=data["title"],
        subtitle="Proposition de Transformation & Business Case MVP (7 Piliers)",
        date_str="Juillet 2026",
        confidentiality="STRICTEMENT CONFIDENTIEL — PROPOSITION PARTENAIRE",
        client_name="SanlamAllianz Group"
    )
    
    # Configure Headers & Footers
    setup_headers_footers(doc, "Fraud Intelligence Accelerator")
    
    # 1. Introduction et Positionnement
    add_heading_1(doc, "1. Positionnement Produit : Fraud Intelligence Platform")
    add_paragraph(doc, "Pour maîtriser son loss ratio et améliorer sa rentabilité technique, SanlamAllianz doit automatiser et fiabiliser la détection des sinistres frauduleux. Au lieu de proposer une boîte à outils algorithmique et complexe pour data scientists, Ekinox propose le déploiement d'une plateforme métier intégrée : Ekinox Fraud Intelligence Accelerator.")
    add_paragraph(doc, "Cette solution combine l'ingestion automatique des rapports d'experts, l'évaluation sémantique, la détection des réseaux de fraude organisée, et la capitalisation immédiate des dossiers résolus. Elle permet de classer instantanément les sinistres et d'accompagner l'enquêteur pas-à-pas avec des alertes entièrement explicitables.")
    
    # 2. Product Blueprint
    add_heading_1(doc, "2. Vision Produit & Principes de Conception")
    add_paragraph(doc, data["product_blueprint"]["vision"])
    add_heading_2(doc, "Principes de Conception du MVP")
    for principle in data["product_blueprint"]["principles"]:
        add_paragraph(doc, principle["description"], bold_prefix=f"• {principle['name']} : ")
        
    # 3. Architecture Technique et Intégration Socle Commun
    add_heading_1(doc, "3. Architecture Technique & Socle Commun")
    add_paragraph(doc, data["technical_architecture"]["overview"])
    add_paragraph(doc, data["technical_architecture"]["integration"], bold_prefix="Intégration SI & Core Claims : ")
    
    add_heading_2(doc, "Composants Technologiques Clés")
    components = data["technical_architecture"]["core_components"].split(". ")
    for comp in components:
        cleaned_comp = comp.strip()
        if not cleaned_comp:
            continue
        # Clean number prefix if present (e.g. "1. " or just "1 ")
        if cleaned_comp[0].isdigit():
            dot_idx = cleaned_comp.find('.')
            if dot_idx != -1 and dot_idx < 3:
                cleaned_comp = cleaned_comp[dot_idx+1:].strip()
            else:
                cleaned_comp = cleaned_comp[1:].strip()
        add_bullet_point(doc, cleaned_comp)
            
    doc.add_page_break()
    
    # 4. Les 7 Capacités Métier de l'Accélérateur V1
    add_heading_1(doc, "4. Cadrage des 7 Capacités du Fraud Intelligence Accelerator")
    add_paragraph(doc, "Chaque pilier de la plateforme associe des bénéfices d'affaires tangibles pour le gestionnaire de sinistres (Section Client) et des spécifications d'architecture applicative pour l'équipe technique (Section Équipe Interne).")
    
    for idx, cap in enumerate(data["capacities"]):
        name = cap["name"]
        if name.startswith(f"{idx+1}."):
            name = name[len(str(idx+1))+2:].strip()
            
        add_heading_2(doc, f"Capacité {idx+1} : {name}")
        
        # Client block
        add_heading_3(doc, "Section Client — Vision Métier & Cas d'Usage")
        add_paragraph(doc, cap["client_section"]["vision"], bold_prefix="Vision & Objectif : ")
        add_paragraph(doc, cap["client_section"]["value_added"], bold_prefix="Valeur Ajoutée Assureur : ")
        add_paragraph(doc, cap["client_section"]["use_case"], bold_prefix="Cas d'Usage Enquêteur : ")
        add_paragraph(doc, cap["client_section"]["user_impact"], bold_prefix="Impact Utilisateur : ")
        
        # Internal block
        add_heading_3(doc, "Section Équipe Interne — Guide d'Implémentation MVP")
        
        internal_text = (
            f"Données nécessaires pour le MVP :\n{cap['internal_team_section']['data_needed']}\n\n"
            f"Stack technique MVP recommandée :\n{cap['internal_team_section']['tech_stack']}\n\n"
            f"APIs & Services à intégrer/développer :\n{cap['internal_team_section']['apis']}\n\n"
            f"Étapes clés de développement :\n{cap['internal_team_section']['key_steps']}"
        )
        add_callout(doc, internal_text, title=f"Fiche Technique Interne - MVP {name}")
        
    doc.add_page_break()
    
    # 5. Roadmap Produit
    add_heading_1(doc, "5. Trajectoire d'Évolution Produit (Roadmap)")
    add_paragraph(doc, "Le déploiement de la solution anti-fraude s'organise en phases successives, démarrant par un ciblage du matériel automobile pour démontrer la valeur rapidement, avant de s'étendre aux sinistres corporels et aux intégrations avancées.")
    
    roadmap_headers = ["Phase", "Objectifs & Périmètre Fonctionnel"]
    roadmap_data = []
    rm = data["roadmap"]
    phases_order = ["phase_1", "phase_2", "phase_3", "phase_4"]
    for phase_key in phases_order:
        if phase_key in rm:
            phase_label = phase_key.upper().replace("_", " ")
            roadmap_data.append([phase_label, rm[phase_key]])
            
    add_styled_table(doc, roadmap_headers, roadmap_data, column_widths=[1.5, 5.0])
    
    # 6. Business Case & ROI
    add_heading_1(doc, "6. Business Case & Modèle de ROI")
    add_paragraph(doc, "L'analyse de ROI démontre le gain direct lié à la prévention des fuites financières sur les sinistres frauduleux et à l'optimisation opérationnelle de la gestion des cas sains.")
    
    bc = data["business_case"]
    
    add_heading_2(doc, "Hypothèses de Cadrage et Impacts Attendus")
    add_paragraph(doc, bc["assumptions"], bold_prefix="Hypothèses initiales : ")
    add_paragraph(doc, bc["impact"], bold_prefix="Impacts sur la sinistralité & OPEX : ")
    add_paragraph(doc, bc["roi_calculation"], bold_prefix="Calcul du ROI Annuel : ")
    
    add_heading_2(doc, "Seuil de Rentabilité (Point Mort)")
    add_paragraph(doc, bc["break_even"])
    
    # Visual Table
    be_headers = ["Indicateur Clé", "Valeur de Référence (SanlamAllianz)", "Cible MVP", "Gains Directs Estimés"]
    be_data = [
        ["Charge Sinistre Globale (Auto)", "100 M€ / an", "98.5 M€ / an (-1.5%)", "1.5 M€ sécurisés / an"],
        ["Faux Positifs en Détection", "30.0%", "<15.0%", "Hausse de productivité enquêteurs"],
        ["Sinistres Simples Auto Matériel", "—", "25% en Fast-Track", "500 k€ d'économies de gestion (OPEX)"],
        ["Valeur Annuelle Créée", "—", "—", "2 000 000 € / an"]
    ]
    add_styled_table(doc, be_headers, be_data, column_widths=[2.4, 1.3, 1.3, 1.5])
    
    # 7. Conditions de Succès et Gouvernance
    add_heading_1(doc, "7. Modèle de Livraison & Gouvernance")
    add_paragraph(doc, "Le cycle d'intégration et de calage des modèles prédictifs du MVP s'étend sur 8 à 10 semaines.")
    
    add_heading_2(doc, "Gouvernance Projet")
    add_bullet_point(doc, "Sprint Planning toutes les 2 semaines et revues conjointes des performances de scoring.")
    add_bullet_point(doc, "Comité technique hebdomadaire pour l'extraction et la mise en conformité RGPD des données historiques.")
    
    add_heading_2(doc, "Conditions de Succès Métier")
    add_bullet_point(doc, "Disponibilité d'un historique qualifié de sinistres (contenant des cas de fraude confirmés et des cas sains) sur les 3 dernières années.")
    add_bullet_point(doc, "Accès aux bases de données des garages partenaires pour identifier les anomalies de facturation.")
    add_bullet_point(doc, "Disponibilité de 2 enquêteurs référents pour calibrer le Fraud Investigation Workspace.")
    
    file_path = "c:\\Projects\\Ekinox IA\\Ekinox_Fraud_Intelligence_Accelerator_Business_Case.docx"
    doc.save(file_path)
    print(f"Fraud Intelligence Accelerator DOCX saved to {file_path}")

def main():
    json_path = "c:\\Projects\\Ekinox IA\\final_proposals.json"
    if not os.path.exists(json_path):
        print(f"Error: JSON file not found at {json_path}")
        return
        
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
        
    build_sales_docx(data["accelerators"]["digital_sales"])
    build_fraud_docx(data["accelerators"]["fraud_intelligence"])
    print("All Word documents generated successfully.")

if __name__ == "__main__":
    main()
